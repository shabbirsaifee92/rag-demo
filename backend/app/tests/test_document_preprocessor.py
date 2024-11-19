import unittest
from ..utils.document_preprocessor import DocumentPreprocessor
import io
import os
from PIL import Image
import numpy as np
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class TestDocumentPreprocessor(unittest.TestCase):
    def setUp(self):
        self.preprocessor = DocumentPreprocessor()
        self.test_data_dir = os.path.join(os.path.dirname(__file__), 'test_data')
        os.makedirs(self.test_data_dir, exist_ok=True)

    def create_test_image(self):
        """Create a test image with some text."""
        img = Image.new('RGB', (300, 100), color='white')
        img_path = os.path.join(self.test_data_dir, 'test_image.png')
        img.save(img_path)
        with open(img_path, 'rb') as f:
            return f.read()

    def create_test_pdf(self):
        """Create a simple PDF with test content."""
        pdf_path = os.path.join(self.test_data_dir, 'test.pdf')
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "Test SOX Compliance Document")
        c.drawString(100, 700, "Section 1: Internal Controls")
        c.drawString(100, 650, "This document outlines key controls.")
        c.save()

        with open(pdf_path, 'rb') as f:
            return f.read()

    def create_test_docx(self):
        """Create a test DOCX file with content."""
        doc = docx.Document()
        doc.add_heading('SOX Compliance Test Document', 0)
        doc.add_paragraph('This is a test paragraph about internal controls.')
        doc.add_paragraph('Section 1: Control Environment')

        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Control ID'
        table.cell(0, 1).text = 'Description'
        table.cell(1, 0).text = 'IC-001'
        table.cell(1, 1).text = 'Access Control'

        docx_path = os.path.join(self.test_data_dir, 'test.docx')
        doc.save(docx_path)

        with open(docx_path, 'rb') as f:
            return f.read()

    def test_chunk_text(self):
        """Test text chunking functionality."""
        test_text = "This is a test document. " * 50
        chunks = self.preprocessor._chunk_text(test_text, chunk_size=100, overlap=20)

        self.assertTrue(len(chunks) > 1)
        for chunk in chunks:
            self.assertLessEqual(len(chunk), 100)

        # Test overlap
        if len(chunks) > 1:
            overlap_text = chunks[0][-20:]
            self.assertTrue(chunks[1].startswith(overlap_text))

    def test_calculate_ocr_confidence(self):
        """Test OCR confidence calculation."""
        good_text = "This is a well-formatted text with proper spacing."
        bad_text = "Th1s 1s p00rly f0rm@tted t3xt w1th b@d ch@r@cters."

        good_confidence = self.preprocessor._calculate_ocr_confidence(good_text)
        bad_confidence = self.preprocessor._calculate_ocr_confidence(bad_text)

        self.assertGreater(good_confidence, bad_confidence)
        self.assertGreaterEqual(good_confidence, 0.0)
        self.assertLessEqual(good_confidence, 1.0)

    def test_is_table_block(self):
        """Test table detection in text blocks."""
        table_text = "| Header 1 | Header 2 |\n| Data 1 | Data 2 |"
        non_table_text = "This is regular text without table structure."

        self.assertTrue(self.preprocessor._is_table_block((0, 0, 0, 0, table_text)))
        self.assertFalse(self.preprocessor._is_table_block((0, 0, 0, 0, non_table_text)))

    def test_process_pdf(self):
        """Test PDF processing."""
        pdf_content = self.create_test_pdf()
        chunks = self.preprocessor._process_pdf(pdf_content)

        self.assertTrue(len(chunks) > 0)
        self.assertTrue(any('SOX Compliance' in chunk['content'] for chunk in chunks))
        self.assertTrue(all(chunk['type'] in ['text', 'table', 'annotation'] for chunk in chunks))

    def test_process_docx(self):
        """Test DOCX processing."""
        docx_content = self.create_test_docx()
        chunks = self.preprocessor._process_docx(docx_content)

        self.assertTrue(len(chunks) > 0)
        self.assertTrue(any('SOX Compliance' in chunk['content'] for chunk in chunks))
        self.assertTrue(any('Control ID' in chunk['content'] for chunk in chunks))

    def test_process_document(self):
        """Test overall document processing."""
        # Test PDF processing
        pdf_content = self.create_test_pdf()
        pdf_chunks = self.preprocessor.process_document(pdf_content, 'test.pdf')
        self.assertTrue(len(pdf_chunks) > 0)

        # Test DOCX processing
        docx_content = self.create_test_docx()
        docx_chunks = self.preprocessor.process_document(docx_content, 'test.docx')
        self.assertTrue(len(docx_chunks) > 0)

        # Test unsupported format
        with self.assertRaises(ValueError):
            self.preprocessor.process_document(b'test content', 'test.txt')

    def test_extract_tables_from_pdf_page(self):
        """Test table extraction from PDF pages."""
        pdf_content = self.create_test_pdf()
        doc = self.preprocessor._process_pdf(pdf_content)

        # Verify table detection and extraction
        tables = []
        for chunk in doc:
            if chunk['type'] == 'table':
                tables.append(chunk['content'])

        # At least basic structure checks
        for table in tables:
            self.assertIsInstance(table, str)
            self.assertTrue(len(table) > 0)

    def tearDown(self):
        """Clean up test files."""
        import shutil
        if os.path.exists(self.test_data_dir):
            shutil.rmtree(self.test_data_dir)

if __name__ == '__main__':
    unittest.main()
